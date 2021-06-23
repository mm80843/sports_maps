import pandas as pd
import numpy as np

import matplotlib.pyplot as plt

import cartopy
import cartopy.crs as ccrs
import cartopy.feature as cfeature
import cartopy.io.img_tiles as cimgt
import geopy.distance
#print(cartopy.__version__)
import re

import seaborn as sns


from geopy import distance

from geopandas import GeoDataFrame
from shapely.geometry import Point

from docx import Document
from docx.shared import Inches

import contextily as ctx
import unidecode



## Support for maps

def add_basemap(ax, zoom, url='http://tile.stamen.com/terrain/tileZ/tileX/tileY.png'):
    xmin, xmax, ymin, ymax = ax.axis()
    print("ax.axis = ",ax.axis())
    basemap, extent = ctx.bounds2img(xmin, ymin, xmax, ymax, zoom=zoom, url=url)
    print("extent = ",extent)
    ax.imshow(basemap, extent=extent, interpolation='bilinear')
    print("xyminmax = ",xmin, xmax, ymin, ymax)
    # restore original x/y limits
    ax.axis((xmin, xmax, ymin, ymax))

def AddTable(doc,df):
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])    
    return 1

## now the good dataset

class dataSets:
    # https://github.com/KitwareMedical/IntersonManager/blob/master/IntersonManager.cpp
    # Interesting doc on the registers
    def __init__(self):
        """Knit
        """ 
        print("Object created")

    def CreateDF(self,pathtoinputs):
        self.ref = pathtoinputs
        self.assetsSecondaire =      pd.read_pickle(pathtoinputs+'inputs/fr-en-adresse-et-geolocalisation-etablissements-premier-et-second-degre.pkl.gzip', compression='gzip')
        self.communes_gps =          pd.read_pickle(pathtoinputs+'inputs/communes_gps.pkl.gzip', compression='gzip')
        self.RES_FichesEquipements = pd.read_pickle(pathtoinputs+'inputs/20180110_RES_FichesEquipements.pkl.gzip', compression='gzip').fillna(0)
        self.CP =                    pd.read_pickle(pathtoinputs+'inputs/CP.pkl.gzip', compression='gzip')
        self.Licenses2015 =          pd.read_pickle(pathtoinputs+'inputs/Licences_2015.pkl.gzip', compression='gzip') 
        
        return [self.assetsSecondaire,self.communes_gps,self.RES_FichesEquipements,self.CP,self.Licenses2015]

    ## Support for lycee - sport 

    def min_distance(self,point, lines):
        NONNULL_DIST = lines.distance(point)
        return NONNULL_DIST.min()

    def min_distance_idx(self,point, lines):
        NONNULL_DIST = lines.distance(point).idxmin()
        return self.gdfRES_VilleCible_MS[self.gdfRES_VilleCible_MS.index == NONNULL_DIST].iloc[0].EquNom

    def min_distance_Complexe(self,point, lines):
        NONNULL_DIST = lines.distance(point).idxmin()
        return self.gdfRES_VilleCible_MS[self.gdfRES_VilleCible_MS.index == NONNULL_DIST].iloc[0].InsNom


    def CreateVille(self,VilleCibleNom):
        
        self.StrVilleCible = str(VilleCibleNom)
        VilleCible = self.RES_FichesEquipements[self.RES_FichesEquipements.ComInsee == VilleCibleNom]
        #print(VilleCible)
        self.VilleCible = VilleCible
        self.assetsSecondaire["Code commune"] = pd.to_numeric(self.assetsSecondaire["Code commune"],errors="coerce")
        self.VilleCible_ecoles = self.assetsSecondaire[self.assetsSecondaire["Code commune"] == int(VilleCibleNom)]
        self.CodeVille = list(set(VilleCible.ComInsee.tolist()))[0]
        #print("CodeVille : "+str(self.CodeVille))
        self.LICENSIES = self.Licenses2015[self.Licenses2015.newcog2 == VilleCibleNom]

        self.Eqts = VilleCible.groupby("EquipementTypeLib").EquNbEquIdentique.sum().to_frame()
        VilleEquipmentsNb = self.Eqts.sort_values(by='EquNbEquIdentique', ascending=False).reset_index()
        self.NBEQ = VilleEquipmentsNb.EquNbEquIdentique.sum()
        #print(NBEQ)
        Eqts = VilleCible.groupby("EquipementTypeLib").EquAnneeService.count().to_frame()
        VilleEquipments = Eqts.sort_values(by='EquAnneeService', ascending=False).reset_index()
        #print(VilleEquipments.EquAnneeService.sum())

        self.StatsEquipementVille = pd.merge(VilleEquipments,VilleEquipmentsNb,on="EquipementTypeLib")
        self.StatsEquipementVille.columns = ["Type d'équipement","Nombre de fiches","Nombre d'équipements"]
        
        # Limite à la métropole pour cartographie
        RES_VilleCible = VilleCible.copy()
        RES_VilleCible.EquipementTypeLib = RES_VilleCible.EquipementTypeLib.str.replace("\x92","'")
        self.RES_VilleCible = RES_VilleCible[(RES_VilleCible['EquGpsX'] < 10.029657627) \
                                                                        & (RES_VilleCible['EquGpsX'] > -5.192734) \
                                                                    & (RES_VilleCible['EquGpsY'] < 51.5803338) \
                                                                    & (RES_VilleCible['EquGpsY'] > 40.88264395) ]

        self.VilleCible_POS = self.communes_gps[self.communes_gps.code_insee == self.CodeVille]

        self.SHAPE_Commune = self.CP[self.CP._CodePoste == self.VilleCible_POS.codes_postaux.values[0]]
        self.SHAPE_Commune.crs = {'init': 'epsg:4326'}


        return self.StrVilleCible,self.VilleCible,self.VilleCible_ecoles, self.CodeVille, self.LICENSIES,
        self.Eqts,VilleEquipmentsNb,NBEQ,VilleEquipments,self.StatsEquipementVille, 
        self.RES_VilleCible, self.VilleCible_POS, self.SHAPE_Commune
    
    
    def CreateMap(self,ZOOM=11):
        if not len(self.StrVilleCible):
            print("No city identified")
        else:
            fig, ax = plt.subplots(figsize=(15,15))

            geometry = [Point(xy) for xy in zip(self.RES_VilleCible.EquGpsX, self.RES_VilleCible.EquGpsY)]
            df = self.RES_VilleCible.drop(['EquGpsX', 'EquGpsY'], axis=1)
            crs = {'init': 'epsg:4326'}
            self.gdfRES_VilleCible = GeoDataFrame(df, crs=crs, geometry=geometry)
            self.gdfRES_VilleCible = self.gdfRES_VilleCible.to_crs(epsg=3857)

            geometry = [Point(xy) for xy in zip(self.VilleCible_ecoles["Coordonnee X"], self.VilleCible_ecoles["Coordonnee Y"])]
            df = self.VilleCible_ecoles.drop(["Coordonnee X", "Coordonnee Y"], axis=1)
            crs = {'init': 'epsg:2154'}
            self.gdfEcoles_Ville = GeoDataFrame(df, crs=crs, geometry=geometry)
            self.gdfEcoles_Ville = self.gdfEcoles_Ville.to_crs(epsg=3857)

            self.RES_VilleCible_MultiSport = self.RES_VilleCible[(self.RES_VilleCible.EquipementTypeLib == "Salle multisports") | (self.RES_VilleCible.EquipementTypeLib == "Plateau EPS/Multisports/city-stades") ]
            self.RES_VilleCible_MultiSport["index"] = self.RES_VilleCible_MultiSport.EquipementId
            self.RES_VilleCible_MultiSport = self.RES_VilleCible_MultiSport.set_index("index")

            geometry = [Point(xy) for xy in zip(self.RES_VilleCible_MultiSport.EquGpsX, self.RES_VilleCible_MultiSport.EquGpsY)]
            df = self.RES_VilleCible_MultiSport.drop(['EquGpsX', 'EquGpsY'], axis=1)
            crs = {'init': 'epsg:4326'}
            self.gdfRES_VilleCible_MS = GeoDataFrame(df, crs=crs, geometry=geometry)
            self.gdfRES_VilleCible_MS = self.gdfRES_VilleCible_MS.to_crs(epsg=3857)



            self.SHAPE_Commune.to_crs(epsg=3857).plot(ax =ax,alpha=0.6,color="black") # quantiles or fisher_jenks #cmap = jet

            self.gdfRES_VilleCible_MS.plot(ax = ax,marker='o', label="Multisport", color="yellow", markersize=16*25,alpha=1);


            color_labels = self.gdfRES_VilleCible['EquipementTypeLib'].unique()
            rgb_values = sns.color_palette("Paired", len(color_labels)) # sets on https://seaborn.pydata.org/tutorial/color_palettes.html
            color_map = dict(zip(color_labels, rgb_values))

            grouped = self.gdfRES_VilleCible.groupby('EquipementTypeLib')
            for key, group in grouped:
                group.plot(ax = ax,marker='o', label=key, color=color_map[key], markersize=5*25,alpha=0.9,linewidth=1,edgecolors= "black");

            grouped = self.gdfEcoles_Ville.groupby('Dénomination principale')
            for key, group in grouped:
                group.plot(ax = ax,marker='X', label=key, color="red", markersize=10*25,alpha=1);

            toner_url = 'http://tile.stamen.com/toner/tileZ/tileX/tileY.png'
            terrain_url='http://tile.stamen.com/terrain/tileZ/tileX/tileY.png'
            # http://tile.stamen.com/toner/{z}/{x}/{y}.png
            add_basemap(ax=ax, zoom=ZOOM, url=toner_url)

            ax.set_axis_off()
            plt.legend(loc='best')
            TITLE  = "Ville: "+self.VilleCible.ComLib.iloc[0]+"\n"+str(len(self.RES_VilleCible)) + " équipements sportifs & "
            TITLE += str(len(self.VilleCible_ecoles))+" batiments scolaires.\n"

            plt.title(TITLE,fontsize = 14)
            plt.tight_layout()

            plt.savefig("outputs/"+str(self.CodeVille)+"_terrain.png")

        
        return fig

    def PrepareDocument(self):

        TABLEAUEQUIPEMENTS = self.gdfRES_VilleCible[["EquipementId","InsNom","EquNom","EquipementTypeLib","EquNbEquIdentique","GestionTypeProprietairePrincLib","EquAnneeService","EquErpCategorie"]]
        TABLEAUEQUIPEMENTS.EquAnneeService = TABLEAUEQUIPEMENTS.EquAnneeService.astype(int)
        TABLEAUEQUIPEMENTSEquipementId = TABLEAUEQUIPEMENTS.EquipementId.astype(int)
        TABLEAUEQUIPEMENTS.EquErpCategorie = TABLEAUEQUIPEMENTS.EquErpCategorie.astype(int)
        TABLEAUEQUIPEMENTS.columns = ['ID', 'Complexe', 'Equipement', 'Type', "Nombre",'Gestionnaire', 'Année de mise en service', 'Categorie ERP']
        TABLEAUEQUIPEMENTS = TABLEAUEQUIPEMENTS.sort_values(['Complexe',"Equipement"],ascending=False)
        self.TABLEAUEQUIPEMENTS = TABLEAUEQUIPEMENTS


        self.gdfEcoles_Ville['distance_to_equpt'] = self.gdfEcoles_Ville.geometry.apply(self.min_distance, args=(self.gdfRES_VilleCible_MS,))
        self.gdfEcoles_Ville['nearest'] = self.gdfEcoles_Ville.geometry.apply(self.min_distance_idx, args=(self.gdfRES_VilleCible_MS,))
        self.gdfEcoles_Ville['nearest_complex'] = self.gdfEcoles_Ville.geometry.apply(self.min_distance_Complexe, args=(self.gdfRES_VilleCible_MS,))


        self.DataEcoles = self.gdfEcoles_Ville[["Appellation officielle","Dénomination principale","distance_to_equpt","nearest","nearest_complex"]]
        self.DataEcoles.distance_to_equpt = self.DataEcoles.distance_to_equpt.astype(int)
        self.DataEcoles.columns = ["Etablissement","Type","Equipement le plus proche (m)","Type","Complexe"]
        #self.DataEcoles.head()

        self.Population = self.LICENSIES.pop_2014.iloc[0]

        StatsClubs = self.LICENSIES.pivot_table(index='nomFede', values='l_2015', aggfunc='sum', fill_value=0)
        StatsClubs = StatsClubs.sort_values(by='l_2015', ascending=False).reset_index()
        StatsClubs = StatsClubs[StatsClubs.l_2015 >0]
        StatsClubs["% population"] = StatsClubs.l_2015 / self.Population
        StatsClubs['% population'] = StatsClubs[['% population']].applymap(lambda x: "{0:.2f} %".format(x*100))
        StatsClubs.columns = [["Nom de la fédération","Licenciés","% de la population communale"]]
        self.StatsClubs = StatsClubs

        self.SOURCES = "Données équipements sportifs: https://www.data.gouv.fr/fr/datasets/recensement-des-equipements-sportifs-espaces-et-sites-de-pratiques/ \n"
        self.SOURCES += "Données éducation: https://www.data.gouv.fr/fr/datasets/adresse-et-geolocalisation-des-etablissements-denseignement-du-premier-et-second-degres-1/ \n"
        self.SOURCES += "Découpage communal: https://www.data.gouv.fr/fr/datasets/decoupage-administratif-communal-francais-issu-d-openstreetmap/ \n"
        self.SOURCES += "Licenses sportives 2015: https://www.data.gouv.fr/fr/datasets/donnees-geocodees-issues-du-recensement-des-licences-et-clubs-aupres-des-federations-sportives-agreees-par-le-ministere-charge-des-sports/ \n"
        

    def createReport(self,template="",output=""):

        # Prepare supporting tables and analyses
        self.PrepareDocument()

        if template == "":
            template = self.ref + "templates/report_template.docx"
        if  output == "":
            output = 'outputs/rapports_'+str(self.CodeVille)+'.docx'

        document = Document(docx=template)

        document.add_heading('Revue de la ville de '+self.VilleCible.ComLib.iloc[0]+' ('+str(self.CodeVille)+")", 0)
        document.add_heading("Sources", level=1)
        document.add_paragraph(self.SOURCES)
        document.add_heading("Liste des équipements de la ville", level=1)

        document.add_paragraph("Il y a "+str(int(self.NBEQ))+" équipements sportifs pour "+str(int(self.Population))+" habitants, soit un ratio de "+str(int(self.NBEQ*10000/self.Population))+" équipements pour 10.000 habitants.")

        AddTable(document,self.StatsEquipementVille)
        document.add_page_break()

        document.add_heading("Vue d'ensemble de la ville", level=1)
        document.add_picture("outputs/"+str(self.CodeVille)+"_terrain.png", width=Inches(6.2)) 
        document.add_page_break()


        document.add_heading('Revue des équipements', level=1)
        AddTable(document,self.TABLEAUEQUIPEMENTS)               
        document.add_page_break()

        document.add_heading('Revue des collèges et lycées', level=1)
        AddTable(document,self.DataEcoles)       

        document.save(output)
        print("Document sauvé sous "+output+ " .")



    def createReportMarkdown(self,template="",output=""):

            if  output == "":
                output = 'outputs/rapports_'+str(self.CodeVille)+'.md'
            # Prepare supporting tables and analyses
            self.PrepareDocument()
            MD = "".encode('utf8').decode('latin1') 
            MD = '# Revue de la ville de '+self.VilleCible.ComLib.iloc[0]+' ('+str(self.CodeVille)+")\n"

            MD += "## Sources\n\n"
            MD += self.SOURCES

            MD += "## Liste des équipements de la ville\n\n"
            MD += "\n\nIl y a "+str(int(self.NBEQ))+" équipements sportifs pour "+str(int(self.Population))+ " habitants, soit un ratio de "+str(int(self.NBEQ*10000/self.Population))+" équipements pour 10.000 habitants."
            MD += "\n\n"+self.StatsEquipementVille.to_markdown()

            MD += "\n\n## Vue d'ensemble de la ville\n\n"
            MD += "![]("+str(self.CodeVille)+"_terrain.png)\n\n"

            MD += "## Revue des équipements \n\n"
            MD += "\n\n"+self.TABLEAUEQUIPEMENTS.to_markdown()

            MD += "\n\n## Revue des collèges et lycées\n\n"
            MD += "\n\n"+self.DataEcoles.to_markdown()    

            #saving the file   
            with open(output, 'w',encoding='latin-1') as f:
                f.write(MD)
            self.MD = MD
            print("Document sauvé sous "+output+ " .")

def remove_accents(input_str):
    nfkd_form = unicodedata.normalize('NFKD', input_str)
    only_ascii = nfkd_form.encode('ASCII', 'ignore')
    return only_ascii